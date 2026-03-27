from docx import Document
import os

def create_resume(filename, name, title, contact, summary, experience, education, skills, projects):
    doc = Document()
    
    # Title
    doc.add_heading(name, 0)
    doc.add_paragraph(f"{title} | {contact}")
    
    # Summary
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(summary)
    
    # Experience
    doc.add_heading('Professional Experience', level=1)
    for exp in experience:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(exp).bold = True
    
    # Skills
    doc.add_heading('Technical Skills', level=1)
    doc.add_paragraph(skills)
    
    # Projects
    doc.add_heading('Key Projects', level=1)
    for proj in projects:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(proj)
        
    # Education
    doc.add_heading('Education', level=1)
    doc.add_paragraph(education)

    # Save
    doc.save(filename)
    print(f"Created {filename}")

# Create output folder
os.makedirs("sample_resumes", exist_ok=True)

# 1. Strong Python Backend Dev
create_resume(
    "sample_resumes/Sarah_Johnson_Python_Backend.docx",
    "Sarah Johnson",
    "Backend Software Engineer",
    "sarah.j@email.com | 555-0101",
    "Experienced Python developer specializing in building scalable REST APIs and database management.",
    ["Backend Dev at TechCorp (2020-Present): Built microservices with Python, Flask, and PostgreSQL.", 
     "Junior Dev at StartupInc (2018-2020): Wrote SQL queries and managed AWS deployments."],
    "B.S. Computer Science, State University",
    "Python, Flask, Django, SQL, PostgreSQL, AWS, Docker, Git, Supabase, Linux",
    ["E-Commerce API: Built a high-traffic Flask API handling 10k requests/min.",
     "Resume Matcher Database: Designed a Supabase PostgreSQL schema for AI data."]
)

# 2. Frontend React Dev
create_resume(
    "sample_resumes/Michael_Chen_React_Frontend.docx",
    "Michael Chen",
    "Frontend Developer",
    "m.chen@email.com | 555-0202",
    "UI/UX focused frontend engineer with a deep passion for modern Javascript frameworks.",
    ["Frontend Engineer at WebSolutions (2021-Present): Lead development of primary React web app.",
     "Intern at DesignAgency (2020-2021): Created responsive UI components using CSS and HTML."],
    "B.A. Graphic Design, Arts College",
    "JavaScript, React, Node.js, HTML5, CSS3, Tailwind, Redux, Figma, TypeScript",
    ["Interactive Dashboard: Built a dynamic dashboard using React and Tailwind.",
     "Real-time Chat: Designed UI for a websocket based chat application in JS."]
)

# 3. AI / Machine Learning Data Scientist
create_resume(
    "sample_resumes/David_Kim_AI_ML.docx",
    "David Kim",
    "Machine Learning Engineer",
    "dkim.ai@email.com | 555-0303",
    "Data scientist focused on NLP, transformers, and deep learning implementations.",
    ["Data Scientist at AI-First (2019-Present): Trained custom NLP models using PyTorch.",
     "Data Analyst at FinTech (2017-2019): Analyzed large datasets with Pandas and Numpy."],
    "M.S. Data Science, Tech University",
    "Python, Machine Learning, Artificial Intelligence, PyTorch, Pandas, Numpy, Sentence-Transformers, SQL",
    ["Semantic Search Engine: Implemented cosine similarity search using sentence-transformers.",
     "Predictive Analytics: Built a ML model to predict customer churn with 92% accuracy."]
)

# 4. Excellent Match for Devops (but weak elsewhere)
create_resume(
    "sample_resumes/Emily_Rodriguez_DevOps.docx",
    "Emily Rodriguez",
    "DevOps & Infrastructure Engineer",
    "erod@email.com | 555-0404",
    "Cloud infrastructure specialist with a focus on CI/CD pipelines and containerization.",
    ["DevOps Lead at CloudNet (2018-Present): Managed Kubernetes clusters and AWS infrastructure.",
     "SysAdmin at LocalHost (2015-2018): Maintained Linux servers and automated deployments."],
    "B.S. Information Technology",
    "AWS, Docker, Kubernetes, Linux, CI/CD, Jenkins, Terraform, Bash, Python",
    ["Cloud Migration: Moved legacy monolith to Docker/K8s microservices on AWS.",
     "Automated Pipeline: Reduced deployment time by 80% using Jenkins and Terraform."]
)

print("\nSuccess! 4 sample resumes have been generated in the 'sample_resumes' folder.")
